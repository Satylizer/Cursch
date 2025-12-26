from typing import Tuple, List, Optional
from aiogram.types import Message
from database.models.models import Document, DocumentType
from database.db import Database
import sqlite3
import io
from docx import Document as DocxDocument 
import PyPDF2
import httpx

class DocumentService:
    def __init__(self, db: Database):
        self.db = db

    async def add_document(self, message: Message, user_id: int) -> Tuple[bool, str]:
        if message.document:
            return await self._process_file_document(message, user_id)
        return self._process_text_input(message.text, user_id) # type: ignore

    async def _process_file_document(self, message: Message, user_id: int) -> Tuple[bool, str]:
        file_ext = message.document.file_name.split('.')[-1].lower()  # type: ignore

        try:
            doc_type = DocumentType(file_ext)
        except ValueError:
            return False, "Недопустимый тип файла"

        file = await message.bot.download(message.document)  # type: ignore
        content = await self._read_file_content(file, doc_type) # type: ignore

        doc = Document(
            name=message.document.file_name,  # type: ignore
            content=content,
            type=doc_type
        )

        try:
            self.db.add_document(doc)

            try:
                async with httpx.AsyncClient() as client:
                    await client.post(
                        "http://localhost:8000/api/v1/add_context",
                        json={"text": doc.content},
                        timeout=30
                    )
            except Exception as e:
                raise RuntimeError(f"RAG sync failed: {e}")

            return True, f"Документ '{doc.name}' добавлен"

        except sqlite3.IntegrityError:
            return False, "Документ уже существует"

    async def _read_file_content(self, file: io.BytesIO, doc_type: DocumentType) -> str:
        try:
            if doc_type == DocumentType.TXT:
                return file.read().decode('utf-8')
            elif doc_type == DocumentType.PDF:
                pdf_reader = PyPDF2.PdfReader(file)
                return "\n".join(page.extract_text() for page in pdf_reader.pages)
            elif doc_type in (DocumentType.DOC, DocumentType.DOCX):
                doc = docx.DocxDocument(file) # type: ignore
                return "\n".join(para.text for para in doc.paragraphs)
            else:
                return f"[Бинарный файл типа {doc_type.value}]"
        except Exception as e:
            return f"[Ошибка чтения файла: {str(e)}]"

    def _process_text_input(self, text: str, user_id: int) -> Tuple[bool, str]:
        try:
            name, doc_type = map(str.strip, text.split(",", 1))
            doc_type = DocumentType(doc_type.lower())
            
            doc = Document(
                name=name,
                content=text,
                type=doc_type
            )
            
            self.db.add_document(doc)
            return True, f"Документ '{name}' добавлен"
        except ValueError:
            return False, "Неверный формат. Используйте: название, тип"
        except sqlite3.IntegrityError:
            return False, "Документ уже существует"

    def get_all_documents(self) -> List[Document]:
        return self.db.get_all_documents()

    def get_document_by_name(self, name: str) -> Optional[Document]:
        try:
            return self.db.get_document_by_name(name)
        except sqlite3.Error as e:
            print(f"Ошибка получения документа: {e}")
            return None

    def delete_document(self, name: str) -> Tuple[bool, str]:
        try:
            if self.db.delete_document_by_name(name):
                return True, f"✅ Документ '{name}' успешно удален"
            return False, "Документ с таким именем не найден"
        except sqlite3.Error as e:
            return False, f"Ошибка базы данных: {str(e)}"